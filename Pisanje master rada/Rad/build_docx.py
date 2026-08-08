# -*- coding: utf-8 -*-
r"""
build_docx.py — саставља мастер рад из изворних .md датотека и стилова ETF шаблона.

Улаз:
  Pisanje master rada/Rad/osnova.docx      (шаблон са стиловима, добијен скриптом priprema_sablona.py)
  Pisanje master rada/Rad/NN-*.md          (поглавља)
  Metrike/Analiza/tabele.md                (табеле резултата)
  Metrike/Slike/*.png                      (графички прикази)

Излаз:
  Pisanje master rada/Master rad - Predrag Pesic.docx

Ознаке у изворним датотекама:
  # / ## / ### / ####          нивои наслова
  [SLIKA: fajl | назив]        слика из директоријума Metrike/Slike
  [SCREENSHOT: опис | назив]   место за снимак екрана који корисник накнадно убацује
  [TABELA: назив]              табела чији редови следе у Markdown облику
  [TABELA-IZ: ознака | назив]  табела преузета из Metrike/Analiza/tabele.md
  [JEDNACINA: израз | број]    нумерисана једначина
  [REF] текст                  ставка списка литературе
  [SKRACENICA] скр | значење   ставка списка скраћеница
  [NAPOMENA-PLACEHOLDER: ...]  видљива напомена о делу који се накнадно попуњава
  ```jezik ... ```             исечак програмског кода

Покретање:
  .venv\Scripts\python.exe "Pisanje master rada\Rad\build_docx.py"
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
RAD = ROOT / "Pisanje master rada" / "Rad"
SLIKE = ROOT / "Metrike" / "Slike"
SNIMCI = ROOT / "Pisanje master rada" / "Slike"
TABELE_MD = ROOT / "Metrike" / "Analiza" / "tabele.md"
OSNOVA = RAD / "osnova.docx"
IZLAZ = ROOT / "Pisanje master rada" / "Master rad - Predrag Pesic.docx"

MAX_SLIKA_CM = 13.0
MAX_SNIMKA_CM = 15.0
CONTENT_CM = 17.0

# ─────────────────────────────────────────────────────────────
# ПОДАЦИ О РАДУ
# ─────────────────────────────────────────────────────────────

NASLOV = "Интерактивна визуелизација графовских алгоритама за проналажење пута са подршком вештачке интелигенције"
PODNASLOV = "Мастер рад"
MENTOR = "проф. др Марија Пунт"
KANDIDAT = "Предраг Пешић 2024/3281"
DATUM = "Београд, август 2026."

# Редослед поглавља централног текста
POGLAVLJA = [
    "01-uvod.md",
    "02-algoritmi.md",
    "03-postojeca-resenja.md",
    "04-arhitektura.md",
    "05-rezimi-rada.md",
    "06-vestacka-inteligencija.md",
    "07-metodologija.md",
    "08-rezultati-algoritmi.md",
    "09-evaluacija-ai-ux.md",
    "10-zakljucak.md",
]
PRILOZI = ["13-prilog-a.md", "14-prilog-b.md", "15-prilog-v.md"]
OZNAKE_PRILOGA = ["А", "Б", "В"]

# ─────────────────────────────────────────────────────────────
# ПОМОЋНЕ ФУНКЦИЈЕ ЗА OOXML
# ─────────────────────────────────────────────────────────────


def ocisti_telo(doc):
    """Уклања сав садржај шаблона, а задржава завршни sectPr."""
    body = doc.element.body
    sectprs = body.findall(qn("w:sectPr"))
    zadrzi = sectprs[-1] if sectprs else None
    for child in list(body):
        if child is not zadrzi:
            body.remove(child)
    return zadrzi


def pokupi_sectpr(putanja):
    """Издваја XML описа прва два пресека из изворног шаблона."""
    import zipfile

    with zipfile.ZipFile(putanja) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    return re.findall(r"<w:sectPr\b.*?</w:sectPr>", xml, re.S)


def dodaj_prelom_sekcije(doc, sectpr_xml):
    """Убацује пасус који носи опис пресека, чиме се затвара претходна секција."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    p = doc.add_paragraph()
    p.style = doc.styles["Osnovni tekst"]
    ppr = p._p.get_or_add_pPr()
    frag = sectpr_xml.replace("<w:sectPr ", f"<w:sectPr {nsdecls('w', 'r')} ", 1)
    ppr.append(parse_xml(frag))
    return p


def dodaj_polje(par, uputstvo):
    """Уграђује Word поље (нпр. садржај или списак слика) у задати пасус."""
    r1 = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc)

    r2 = par.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = uputstvo
    r2._r.append(it)

    r3 = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc)

    r4 = par.add_run("Притисните F9 да бисте освежили ово поље.")
    r4.italic = True

    r5 = par.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "end")
    r5._r.append(fc)


def prelom_strane(doc):
    from docx.enum.text import WD_BREAK

    p = doc.add_paragraph()
    p.style = doc.styles["Osnovni tekst"]
    p.add_run().add_break(WD_BREAK.PAGE)


def obelezi_celiju_zaglavlja(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "EDEDED")
    tcpr.append(shd)


def ponovi_zaglavlje(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    trpr.append(el)


def ne_prelamaj_red(row):
    """Спречава да се један ред табеле преломи преко две стране."""
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def zadrzi_uz_sledeci(par):
    """Везује пасус за оно што следи, чиме остаје уз слику или табелу."""
    ppr = par._p.get_or_add_pPr()
    ppr.append(OxmlElement("w:keepNext"))


# ─────────────────────────────────────────────────────────────
# ИНЛАЈН ФОРМАТИРАЊЕ
# ─────────────────────────────────────────────────────────────

# Курзив се препознаје само када звездице стоје уз реч и одвојене су празнином или
# интерпункцијом, чиме се спречава да ознака алгоритма A* буде схваћена као почетак курзива.
INLINE = re.compile(r"(\*\*.+?\*\*|(?<![\w*])\*(?!\s)[^*]+?(?<!\s)\*(?![\w*])|`[^`]+?`)")

# Индекси и експоненти у једначинама: x_{ab}, x^{ab}, x_a, x^a
IZRAZ = re.compile(r"([_^])(?:\{([^}]*)\}|(\w))")


def upisi_izraz(par, izraz):
    """Уписује математички израз уз обраду индекса и експонената."""
    pozicija = 0
    for m in IZRAZ.finditer(izraz):
        if m.start() > pozicija:
            r = par.add_run(izraz[pozicija:m.start()])
            r.italic = True
        r = par.add_run(m.group(2) if m.group(2) is not None else m.group(3))
        r.italic = True
        if m.group(1) == "_":
            r.font.subscript = True
        else:
            r.font.superscript = True
        pozicija = m.end()
    if pozicija < len(izraz):
        r = par.add_run(izraz[pozicija:])
        r.italic = True


def upisi_tekst(par, tekst, osnovna_velicina=None):
    """Уписује текст у пасус уз обраду подебљања, курзива и означавања кода."""
    for deo in INLINE.split(tekst):
        if not deo:
            continue
        if deo.startswith("**") and deo.endswith("**") and len(deo) > 4:
            r = par.add_run(deo[2:-2])
            r.bold = True
        elif deo.startswith("*") and deo.endswith("*") and len(deo) > 2:
            r = par.add_run(deo[1:-1])
            r.italic = True
        elif deo.startswith("`") and deo.endswith("`") and len(deo) > 2:
            r = par.add_run(deo[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r = par.add_run(deo)
        if osnovna_velicina is not None:
            r.font.size = osnovna_velicina


# ─────────────────────────────────────────────────────────────
# УЧИТАВАЊЕ ТАБЕЛА ИЗ ОБРАДЕ МЕРЕЊА
# ─────────────────────────────────────────────────────────────


def ucitaj_spoljne_tabele():
    if not TABELE_MD.exists():
        return {}
    tekst = TABELE_MD.read_text(encoding="utf-8")
    rezultat = {}
    trenutni = None
    redovi = []
    for linija in tekst.splitlines():
        if linija.startswith("### "):
            if trenutni and redovi:
                rezultat[trenutni] = redovi
            trenutni = linija[4:].strip()
            redovi = []
        elif linija.startswith("|") and trenutni:
            redovi.append(linija)
    if trenutni and redovi:
        rezultat[trenutni] = redovi
    return rezultat


SPOLJNE_TABELE = ucitaj_spoljne_tabele()


def nadji_spoljnu_tabelu(oznaka):
    oznaka = oznaka.strip()
    for kljuc, redovi in SPOLJNE_TABELE.items():
        if kljuc.startswith(oznaka):
            return redovi
    for kljuc, redovi in SPOLJNE_TABELE.items():
        if oznaka in kljuc:
            return redovi
    return None


# ─────────────────────────────────────────────────────────────
# ГРАДИТЕЉ ДОКУМЕНТА
# ─────────────────────────────────────────────────────────────


class Graditelj:
    def __init__(self, doc):
        self.doc = doc
        self.poglavlje = 0
        self.potpoglavlje = 0
        self.br_slike = 0
        self.br_tabele = 0
        self.prilog = False
        self.oznaka_priloga = ""
        self.popis_slika = []
        self.popis_tabela = []
        self.upozorenja = []

    # ── наслови ────────────────────────────────────────────

    def naslov(self, nivo, tekst):
        if self.prilog:
            stil = {1: "Prilog - I nivo naslova", 2: "Prilog - II nivo naslova",
                    3: "Prilog - III nivo naslova", 4: "Prilog - IV nivo naslova"}[nivo]
        else:
            stil = {1: "I nivo naslova - Poglavlje", 2: "II nivo naslova - Potpoglavlje",
                    3: "III nivo naslova - Odeljak", 4: "IV nivo naslova - Pododeljak"}[nivo]
        p = self.doc.add_paragraph(style=self.doc.styles[stil])
        if self.prilog:
            # Аутоматска нумерација шаблона даје латинична слова и не
            # почиње изнова у сваком прилогу, па се ћирилична ознака
            # уноси у сам текст наслова.
            self._ukloni_numeraciju(p)
            if nivo == 1:
                p.add_run(f"ПРИЛОГ {self.oznaka_priloga}: {tekst}")
            elif nivo == 2:
                p.add_run(f"{self.oznaka_priloga}.{self.potpoglavlje + 1}. {tekst}")
            else:
                p.add_run(tekst)
        else:
            p.add_run(tekst)
        if nivo == 1:
            self.poglavlje += 1
            self.potpoglavlje = 0
            self.br_slike = 0
            self.br_tabele = 0
        elif nivo == 2:
            self.potpoglavlje += 1
            self.br_slike = 0
            self.br_tabele = 0
        return p

    @staticmethod
    def _ukloni_numeraciju(par):
        ppr = par._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), "0")
        numpr.append(ilvl)
        numpr.append(numid)
        ppr.append(numpr)

    def _prefiks(self):
        if self.prilog:
            osnova = f"{self.oznaka_priloga}.{self.potpoglavlje}" if self.potpoglavlje else self.oznaka_priloga
        else:
            osnova = f"{self.poglavlje}.{self.potpoglavlje}" if self.potpoglavlje else f"{self.poglavlje}"
        return osnova

    # ── основни текст ──────────────────────────────────────

    def pasus(self, tekst):
        p = self.doc.add_paragraph(style=self.doc.styles["Osnovni tekst"])
        upisi_tekst(p, tekst)
        return p

    def nabrajanje(self, tekst, numerisano=False):
        stil = "Numerisano nabrajanje" if numerisano else "Nabrajanje"
        p = self.doc.add_paragraph(style=self.doc.styles[stil])
        upisi_tekst(p, tekst)
        return p

    def jednacina(self, izraz, broj):
        p = self.doc.add_paragraph(style=self.doc.styles["Jednacine"])
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.tab_stops.clear_all()
        pf.tab_stops.add_tab_stop(Cm(CONTENT_CM / 2), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(CONTENT_CM), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        upisi_izraz(p, izraz)
        p.add_run("\t")
        p.add_run(f"({broj})")

    def kod(self, linije):
        for linija in linije:
            p = self.doc.add_paragraph()
            p.style = self.doc.styles["Osnovni tekst"]
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0.6)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            r = p.add_run(linija if linija else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            rpr = r._r.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.append(rf)
            rf.set(qn("w:ascii"), "Consolas")
            rf.set(qn("w:hAnsi"), "Consolas")
            rf.set(qn("w:cs"), "Consolas")

    def napomena(self, tekst):
        p = self.doc.add_paragraph(style=self.doc.styles["Osnovni tekst"])
        r = p.add_run("НАПОМЕНА ЗА ДОПУНУ: ")
        r.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r2 = p.add_run(tekst)
        r2.italic = True
        r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # ── слике ──────────────────────────────────────────────

    def slika(self, ime_datoteke, naziv):
        # Дијаграми и графици стоје уз обраду мерења, снимци екрана уз рад.
        putanja = SLIKE / ime_datoteke
        sirina = MAX_SLIKA_CM
        if not putanja.exists() and (SNIMCI / ime_datoteke).exists():
            putanja = SNIMCI / ime_datoteke
            sirina = MAX_SNIMKA_CM
        self.br_slike += 1
        oznaka = f"{self._prefiks()}.{self.br_slike}"
        p = self.doc.add_paragraph(style=self.doc.styles["Slike/Tabele"])
        zadrzi_uz_sledeci(p)
        if putanja.exists():
            p.add_run().add_picture(str(putanja), width=Cm(sirina))
        else:
            self.upozorenja.append(f"недостаје слика: {ime_datoteke}")
            r = p.add_run(f"[ НЕДОСТАЈЕ ДАТОТЕКА: {ime_datoteke} ]")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        self._opis_slike(oznaka, naziv)

    def mesto_za_snimak(self, opis, naziv):
        self.br_slike += 1
        oznaka = f"{self._prefiks()}.{self.br_slike}"
        p = self.doc.add_paragraph(style=self.doc.styles["Slike/Tabele"])
        zadrzi_uz_sledeci(p)
        r = p.add_run(f"[ МЕСТО ЗА СНИМАК ЕКРАНА: {opis} ]")
        r.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x60, 0xC0)
        self._opis_slike(oznaka, naziv)
        self.upozorenja.append(f"снимак екрана {oznaka}: {opis}")

    def _opis_slike(self, oznaka, naziv):
        p = self.doc.add_paragraph(style=self.doc.styles["Oznaka slike"])
        p.add_run(f"Слика {oznaka}. {naziv}")
        self.popis_slika.append((oznaka, naziv))

    # ── табеле ─────────────────────────────────────────────

    def tabela(self, naziv, redovi_md):
        self.br_tabele += 1
        oznaka = f"{self._prefiks()}.{self.br_tabele}"
        p = self.doc.add_paragraph(style=self.doc.styles["Oznaka tabele"])
        p.add_run(f"Табела {oznaka}. {naziv}")
        zadrzi_uz_sledeci(p)
        self.popis_tabela.append((oznaka, naziv))

        podaci = []
        for linija in redovi_md:
            if re.fullmatch(r"\|[\s:\-|]+\|", linija.strip()):
                continue
            celije = [c.strip() for c in linija.strip().strip("|").split("|")]
            podaci.append(celije)
        if not podaci:
            return
        kolona = max(len(r) for r in podaci)
        podaci = [r + [""] * (kolona - len(r)) for r in podaci]

        t = self.doc.add_table(rows=len(podaci), cols=kolona)
        t.style = self.doc.styles["Table Grid"]
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True

        velicina = Pt(10) if kolona <= 6 else Pt(8.5) if kolona <= 9 else Pt(7.5)
        for i, red in enumerate(podaci):
            for j, vrednost in enumerate(red):
                if i > 0 and re.fullmatch(r"[-+]?\d+\.\d+", vrednost):
                    vrednost = vrednost.replace(".", ",")
                cell = t.cell(i, j)
                cell.text = ""
                par = cell.paragraphs[0]
                par.style = self.doc.styles["Normal"]
                pf = par.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                pf.line_spacing = 1.0
                if i == 0:
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j > 0 and re.fullmatch(r"[-+]?[\d\s.,%]*", vrednost or ""):
                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                upisi_tekst(par, vrednost, velicina)
                if i == 0:
                    for r in par.runs:
                        r.bold = True
                    obelezi_celiju_zaglavlja(cell)
        ponovi_zaglavlje(t.rows[0])
        for red in t.rows:
            ne_prelamaj_red(red)


# ─────────────────────────────────────────────────────────────
# ОБРАДА ЈЕДНЕ ИЗВОРНЕ ДАТОТЕКЕ
# ─────────────────────────────────────────────────────────────

RE_SLIKA = re.compile(r"^\[SLIKA:\s*(.+?)\s*\|\s*(.+?)\s*\]$")
RE_SNIMAK = re.compile(r"^\[SCREENSHOT:\s*(.+?)\s*\|\s*(.+?)\s*\]$")
RE_TABELA = re.compile(r"^\[TABELA:\s*(.+?)\s*\]$")
RE_TABELA_IZ = re.compile(r"^\[TABELA-IZ:\s*(.+?)\s*\|\s*(.+?)\s*\]$")
RE_JEDNACINA = re.compile(r"^\[JEDNACINA:\s*(.+?)\s*\|\s*(.+?)\s*\]$")
RE_NAPOMENA = re.compile(r"^\[NAPOMENA-PLACEHOLDER:\s*(.+?)\s*\]$", re.S)


def obradi_datoteku(g: Graditelj, putanja: Path):
    linije = putanja.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(linije)
    while i < n:
        red = linije[i]
        strip = red.strip()

        if not strip:
            i += 1
            continue

        # исечак кода
        if strip.startswith("```"):
            i += 1
            blok = []
            while i < n and not linije[i].strip().startswith("```"):
                blok.append(linije[i])
                i += 1
            i += 1
            g.kod(blok)
            continue

        # наслови
        m = re.match(r"^(#{1,4})\s+(.*)$", strip)
        if m:
            g.naslov(len(m.group(1)), m.group(2).strip())
            i += 1
            continue

        m = RE_JEDNACINA.match(strip)
        if m:
            g.jednacina(m.group(1), m.group(2))
            i += 1
            continue

        m = RE_SLIKA.match(strip)
        if m:
            g.slika(m.group(1), m.group(2))
            i += 1
            continue

        m = RE_SNIMAK.match(strip)
        if m:
            g.mesto_za_snimak(m.group(1), m.group(2))
            i += 1
            continue

        m = RE_TABELA_IZ.match(strip)
        if m:
            redovi = nadji_spoljnu_tabelu(m.group(1))
            if redovi is None:
                g.upozorenja.append(f"није пронађена спољна табела: {m.group(1)}")
                g.pasus(f"[ НИЈЕ ПРОНАЂЕНА ТАБЕЛА: {m.group(1)} ]")
            else:
                g.tabela(m.group(2), redovi)
            i += 1
            continue

        m = RE_TABELA.match(strip)
        if m:
            naziv = m.group(1)
            i += 1
            while i < n and not linije[i].strip().startswith("|"):
                i += 1
            redovi = []
            while i < n and linije[i].strip().startswith("|"):
                redovi.append(linije[i])
                i += 1
            g.tabela(naziv, redovi)
            continue

        m = RE_NAPOMENA.match(strip)
        if m:
            g.napomena(m.group(1))
            i += 1
            continue

        if strip.startswith("[REF] "):
            p = g.doc.add_paragraph(style=g.doc.styles["Referenca"])
            upisi_tekst(p, strip[6:])
            i += 1
            continue

        if strip.startswith("[SKRACENICA] "):
            i += 1
            continue

        if re.match(r"^\d+\.\s+", strip):
            g.nabrajanje(re.sub(r"^\d+\.\s+", "", strip), numerisano=True)
            i += 1
            continue

        if strip.startswith("- ") or strip.startswith("* "):
            g.nabrajanje(strip[2:], numerisano=False)
            i += 1
            continue

        g.pasus(strip)
        i += 1


def tabela_skracenica(g: Graditelj, putanja: Path):
    stavke = []
    for linija in putanja.read_text(encoding="utf-8").splitlines():
        if linija.startswith("[SKRACENICA] "):
            skr, znac = linija[13:].split("|", 1)
            stavke.append((skr.strip(), znac.strip()))
    t = g.doc.add_table(rows=len(stavke), cols=2)
    t.style = g.doc.styles["Table Grid"]
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (skr, znac) in enumerate(stavke):
        for j, vrednost in enumerate((skr, znac)):
            par = t.cell(i, j).paragraphs[0]
            par.style = g.doc.styles["Normal"]
            par.paragraph_format.space_before = Pt(1)
            par.paragraph_format.space_after = Pt(1)
            r = par.add_run(vrednost)
            r.font.size = Pt(11)
            if j == 0:
                r.bold = True
    t.columns[0].width = Cm(3.5)
    t.columns[1].width = Cm(12.0)


# ─────────────────────────────────────────────────────────────
# НАСЛОВНА СТРАНА
# ─────────────────────────────────────────────────────────────


def naslovna_strana(doc):
    for tekst in ("Универзитет у Београду", "Електротехнички факултет"):
        p = doc.add_paragraph(style=doc.styles["Zaglavlje naslovne strane"])
        p.add_run(tekst)

    p = doc.add_paragraph(style=doc.styles["Slike/Tabele"])
    logo = next((d / "etf-logo.png" for d in (SNIMCI, SLIKE) if (d / "etf-logo.png").exists()), None)
    if logo:
        p.add_run().add_picture(str(logo), height=Cm(3.2))
    else:
        r = p.add_run("[ МЕСТО ЗА ЛОГО ФАКУЛТЕТА ]")
        r.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x60, 0xC0)

    for _ in range(5):
        doc.add_paragraph(style=doc.styles["Osnovni tekst"])

    p = doc.add_paragraph(style=doc.styles["Naslov teze"])
    p.add_run(NASLOV)

    p = doc.add_paragraph(style=doc.styles["Podnaslov teze"])
    p.add_run(PODNASLOV)

    for _ in range(4):
        doc.add_paragraph(style=doc.styles["Osnovni tekst"])

    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    podaci = [("Ментор:", "Кандидат:"), (MENTOR, KANDIDAT)]
    for i, red in enumerate(podaci):
        for j, vrednost in enumerate(red):
            par = t.cell(i, j).paragraphs[0]
            try:
                par.style = doc.styles["Naslovna kandidat i mentor"]
            except KeyError:
                par.style = doc.styles["Normal"]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            par.add_run(vrednost)
    t.columns[0].width = Cm(8.5)
    t.columns[1].width = Cm(8.5)

    for _ in range(3):
        doc.add_paragraph(style=doc.styles["Osnovni tekst"])

    p = doc.add_paragraph(style=doc.styles["Vreme predaje teze"])
    p.add_run(DATUM)


# ─────────────────────────────────────────────────────────────
# ГЛАВНИ ТОК
# ─────────────────────────────────────────────────────────────


def zavrsna_celina(doc, naslov):
    """Започиње нову страну и уписује наслов целине иза централног текста."""
    prelom_strane(doc)
    p = doc.add_paragraph(style=doc.styles["Sadržaj/Literatura"])
    p.add_run(naslov)
    return p


def main():
    if not OSNOVA.exists():
        sys.exit("Не постоји osnova.docx. Прво покрените priprema_sablona.py.")

    sectpr_lista = pokupi_sectpr(OSNOVA)
    doc = Document(str(OSNOVA))
    ocisti_telo(doc)

    naslovna_strana(doc)
    dodaj_prelom_sekcije(doc, sectpr_lista[0])

    g = Graditelj(doc)

    # Захвалница и сажетак иду између насловне стране и садржаја.
    p = doc.add_paragraph(style=doc.styles["Sadržaj/Literatura"])
    p.add_run("Захвалница")
    obradi_datoteku(g, RAD / "00-zahvalnica.md")

    prelom_strane(doc)
    p = doc.add_paragraph(style=doc.styles["Sadržaj/Literatura"])
    p.add_run("Сажетак")
    obradi_datoteku(g, RAD / "00-sazetak.md")

    prelom_strane(doc)
    p = doc.add_paragraph(style=doc.styles["Sadržaj/Literatura"])
    p.add_run("Садржај")
    p = doc.add_paragraph(style=doc.styles["Osnovni tekst"])
    dodaj_polje(p, 'TOC \\o "1-3" \\t "Sadržaj/Literatura,1,Prilog - I nivo naslova,1,'
                   'Prilog - II nivo naslova,2" \\h \\z \\u ')

    dodaj_prelom_sekcije(doc, sectpr_lista[1])

    for ime in POGLAVLJA:
        obradi_datoteku(g, RAD / ime)

    # Литература
    zavrsna_celina(doc, "Литература")
    obradi_datoteku(g, RAD / "11-literatura.md")

    # Списак скраћеница
    zavrsna_celina(doc, "Списак скраћеница")
    tabela_skracenica(g, RAD / "12-skracenice.md")

    # Списак слика
    zavrsna_celina(doc, "Списак слика")
    p = doc.add_paragraph(style=doc.styles["Osnovni tekst"])
    dodaj_polje(p, r'TOC \h \z \t "Oznaka slike" \c ')

    # Списак табела
    zavrsna_celina(doc, "Списак табела")
    p = doc.add_paragraph(style=doc.styles["Osnovni tekst"])
    dodaj_polje(p, r'TOC \h \z \t "Oznaka tabele" \c ')

    # Прилози
    g.prilog = True
    for oznaka, ime in zip(OZNAKE_PRILOGA, PRILOZI):
        g.oznaka_priloga = oznaka
        g.potpoglavlje = 0
        g.br_slike = 0
        g.br_tabele = 0
        obradi_datoteku(g, RAD / ime)

    cp = doc.core_properties
    cp.author = "Предраг Пешић"
    cp.last_modified_by = "Предраг Пешић"
    cp.title = NASLOV
    cp.subject = PODNASLOV
    cp.category = ""
    cp.comments = ""
    cp.keywords = ""

    doc.save(str(IZLAZ))

    print(f"Сачувано: {IZLAZ}")
    print(f"Слика: {len(g.popis_slika)}   Табела: {len(g.popis_tabela)}")
    if g.upozorenja:
        print("\nСтавке које траже допуну:")
        for u in g.upozorenja:
            print(f"  - {u}")

    izvestaj = RAD / "popis-slika-i-tabela.md"
    with open(izvestaj, "w", encoding="utf-8") as fh:
        fh.write("# Попис слика\n\n")
        for oznaka, naziv in g.popis_slika:
            fh.write(f"- Слика {oznaka}. {naziv}\n")
        fh.write("\n# Попис табела\n\n")
        for oznaka, naziv in g.popis_tabela:
            fh.write(f"- Табела {oznaka}. {naziv}\n")
    print(f"Попис: {izvestaj}")


if __name__ == "__main__":
    main()
